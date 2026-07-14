from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PHP = "PHP "
RATE = 58
AS_OF_DATE = "June 21, 2026"


def php(usd):
    return int(round(usd * RATE))


def money_php(amount):
    return f"{PHP}{amount:,.0f}"


domain_note = (
    "Domain name cost. This is annual, not monthly. If Axon keeps Namecheap, PremiumDNS is recommended "
    "for a business-critical domain because it adds a 100% DNS uptime SLA, global Anycast DNS, DNSSEC support, "
    "DNS-layer DDoS protection, ALIAS record support, and up to 2M DNS queries/month. Hostinger should be considered "
    "only if the CEO wants registrar consolidation, simpler billing, support ownership, or bundled DNS/domain management."
)

runtime = [
    ("Vercel", f"USD 30 ({money_php(php(30))})/month", php(30), "Paid and active. USD 20 Pro plan plus USD 10 Speed Insights. Hosts the website and backend API."),
    ("Supabase PostgreSQL", f"USD 25 ({money_php(php(25))})/month", php(25), "Paid and active. Production database for users, events, tickets, orders, registrations, and admin records."),
    ("Supabase UAT compute", f"USD 10 ({money_php(php(10))})/month", php(10), "Paid and active compute allocation for the separate Axon UAT environment."),
    ("Upstash Redis", "Free", php(20), "Fast temporary data for reservations, rate limiting, login protection, and verification flows."),
    ("Cloudinary", "Free", php(89), "Stores and delivers event images and uploaded proof/payment images."),
    ("hCaptcha", "Free", php(99), "Bot and abuse protection for login, registration, and checkout-related forms."),
    ("Google OAuth", "Free", 0, "Optional Google login. Currently hidden due to a bug; existing login works."),
    ("Google Maps", "Free", php(275), "Future reserve only. Used for event location display/linking; free usage may be enough initially."),
    ("Mapbox", "Free", None, "Pay-as-you-go. Used for static location images; excluded from fixed totals."),
    ("Brevo SMTP", "Free", php(18), "Sends OTP codes, ticket confirmations, QR emails, and admin resend emails."),
    ("Sentry", "Free", php(26), "Production error tracking so issues can be fixed before repeated customer reports."),
    ("Domain", "Namecheap: PHP 351/year (domain + PremiumDNS)", "Hostinger: PHP 2,777 first year; PHP 637/year renewal", domain_note),
]

ai_tools = [
    ("GitHub Copilot", money_php(php(39)), "Discontinue", "Remove because it overlaps with ChatGPT/Codex and Claude Code."),
    ("ChatGPT / Codex", money_php(php(20)), "Pro: PHP 6,490/month", "Commercialization upgrade: higher Codex usage, deeper research, larger context, and sustained release support."),
    ("Claude Code", money_php(php(20)), "Max: USD 100 / PHP 5,800/month", "Commercialization upgrade: higher limits and priority access for Claude Code, UAT remediation, and production incident work."),
]

current_app_runtime_usd = 30 + 25 + 10
current_app_runtime = php(current_app_runtime_usd)
current_ai_usd = 39 + 20 + 20
current_ai = php(current_ai_usd)
current_committed_usd = current_app_runtime_usd + current_ai_usd
current_committed_monthly = current_app_runtime + current_ai
historical_may_ai_usd = 39
historical_june_ai_usd = 39 + 20 + 20
historical_june_runtime_usd = 30 + 25 + 10
historical_total_usd = historical_may_ai_usd + historical_june_ai_usd + historical_june_runtime_usd
historical_total = php(historical_total_usd)
runtime_base = sum(x[2] for x in runtime if isinstance(x[2], int) and x[0] != "Google Maps")
runtime_conservative = runtime_base + php(275)
ai_future = 6490 + php(100)
current_domain_first_year = php(1.18)
namecheap_premium_dns_first_year = php(4.88)
namecheap_recommended_first_year = current_domain_first_year + namecheap_premium_dns_first_year
hostinger_domain_first_year = 2777
hostinger_domain_renewal = 637
future_base_monthly = runtime_base + ai_future
future_cons_monthly = runtime_conservative + ai_future


COLORS = {
    "navy": RGBColor(11, 37, 69),
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "muted": RGBColor(92, 104, 116),
    "light": "F2F4F7",
    "blue_fill": "E8EEF5",
    "green_fill": "EAF5EF",
    "gold_fill": "FFF4D6",
    "white": RGBColor(255, 255, 255),
    "black": RGBColor(0, 0, 0),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        elem = mar.find(qn(f"w:{m}"))
        if elem is None:
            elem = OxmlElement(f"w:{m}")
            mar.append(elem)
        elem.set(qn("w:w"), str(v))
        elem.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def paragraph_border_bottom(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", size=11, bold=False, color=None, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, COLORS["blue"], 16, 8
    elif level == 2:
        size, color, before, after = 13, COLORS["blue"], 12, 6
    else:
        size, color, before, after = 12, COLORS["dark_blue"], 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run(r, size=size, bold=True, color=color)
    return p


def add_callout(doc, title, body, fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9120])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="C7D7EA", size="4")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, size=11.5, bold=True, color=COLORS["navy"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=COLORS["black"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def fill_cell(cell, text, size=9.3, bold=False, color=None, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    for p in cell.paragraphs:
        p.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, header_fill="F2F4F7", font_size=8.9, right_cols=(1, 2)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_shading(c, header_fill)
        set_cell_border(c, color="C9D3DF", size="6")
        fill_cell(c, h, size=8.7, bold=True, color=COLORS["navy"], align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            set_cell_border(c, color="DDE5EE", size="4")
            align = WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else None
            fill_cell(c, str(value), size=font_size, color=COLORS["black"], align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_kpi_table(doc):
    headers = ["Scenario", "Monthly Expense", "Annual Expense", "Notes"]
    rows = [
        ("Current app runtime", f"USD {current_app_runtime_usd} / {money_php(current_app_runtime)}", money_php(current_app_runtime * 12 + namecheap_recommended_first_year), "Vercel USD 30, Supabase Pro USD 25, and Supabase UAT compute USD 10 are paid and active. Annual total also includes Namecheap domain + PremiumDNS."),
        ("Total current committed cost", f"USD {current_committed_usd} / {money_php(current_committed_monthly)}", money_php(current_committed_monthly * 12 + namecheap_recommended_first_year), "Current app runtime plus GitHub Copilot and ChatGPT/Codex. This is the present run-rate before other commercial upgrades."),
        ("Recommended commercial baseline", money_php(future_base_monthly), money_php(future_base_monthly * 12 + hostinger_domain_first_year), "Projected production budget. Includes Vercel with Speed Insights and Supabase; excludes Google Maps paid upgrade and pay-as-you-go items."),
        ("Conservative reserve scenario", money_php(future_cons_monthly), money_php(future_cons_monthly * 12 + hostinger_domain_first_year), "Adds Google Maps paid upgrade as a reserve scenario."),
    ]
    add_table(doc, headers, rows, [1900, 1850, 1650, 3960], header_fill="E8EEF5", font_size=8.4)


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3000, 3120, 3000])
    items = [
        ("Live App Runtime", money_php(current_app_runtime), "USD 65/month: Vercel + Supabase Prod/UAT"),
        ("Total Committed Now", money_php(current_committed_monthly), f"USD {current_committed_usd}/month including AI tools"),
        ("Commercial Baseline", money_php(future_base_monthly), "projected monthly operating budget"),
    ]
    for idx, (label, value, note) in enumerate(items):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "EAF5EF" if idx < 2 else "FFF4D6")
        set_cell_border(cell, color="C9D3DF", size="6")
        set_cell_margins(cell, top=180, bottom=180, start=160, end=160)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label)
        set_run(r, size=8.5, bold=True, color=COLORS["muted"])
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(value)
        set_run(r2, size=14, bold=True, color=COLORS["navy"])
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(note)
        set_run(r3, size=8, color=COLORS["muted"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_historical_expense_table(doc):
    rows = [
        ("May 2026", "GitHub Copilot", "USD 39", money_php(php(39)), "Paid"),
        ("June 2026", "Claude Code", "USD 20", money_php(php(20)), "Paid"),
        ("June 2026", "ChatGPT / Codex", "USD 20", money_php(php(20)), "Paid"),
        ("June 2026", "GitHub Copilot", "USD 39", money_php(php(39)), "Paid"),
        ("June 2026", "Vercel Pro + Speed Insights", "USD 30", money_php(php(30)), "Paid"),
        ("June 2026", "Supabase Pro", "USD 25", money_php(php(25)), "Paid"),
        ("June 2026", "Supabase UAT compute allocation", "USD 10", money_php(php(10)), "Paid"),
        ("TOTAL", "Recorded software expenses", f"USD {historical_total_usd}", money_php(historical_total), "Paid history"),
    ]
    add_table(doc, ["Billing Month", "Subscription", "Vendor Charge", "PHP Equivalent", "Status"], rows, [1500, 2350, 1600, 1750, 1760], header_fill="E8EEF5", font_size=8.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.10
        r = p.add_run("• ")
        set_run(r, size=10.3, color=COLORS["blue"], bold=True)
        r2 = p.add_run(item)
        set_run(r2, size=10.3)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    for style_name in ("Normal",):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    # First-page masthead.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("AXON TICKETS")
    set_run(r, size=9.5, bold=True, color=COLORS["blue"])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("CEO Cost Summary")
    set_run(tr, size=25, bold=True, color=COLORS["navy"])

    subtitle = add_para(
        doc,
        "Costs paid, monthly costs today, and the potential budget for commercialization",
        size=12,
        color=COLORS["muted"],
        after=8,
    )
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="2E74B5", size="12")
    rule.paragraph_format.space_after = Pt(10)

    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [1800, 7320])
    meta_rows = [
        ("Cost position", f"As of {AS_OF_DATE}"),
        ("Planning rate", f"USD 1 = PHP {RATE}"),
    ]
    for row, (label, value) in zip(meta.rows, meta_rows):
        for cell in row.cells:
            set_cell_border(cell, color="FFFFFF", size="0")
            set_cell_margins(cell, top=45, bottom=45, start=80, end=80)
        fill_cell(row.cells[0], label, size=9.3, bold=True, color=COLORS["muted"])
        fill_cell(row.cells[1], value, size=9.3, color=COLORS["black"])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3000, 3120, 3000])
    summary_items = [
        ("Paid Since May", money_php(historical_total), "USD 183 actually purchased"),
        ("Monthly Cost Today", money_php(current_committed_monthly), "USD 144 current recurring cost"),
        ("Commercial Monthly Budget", money_php(future_base_monthly), "potential funded operating cost"),
    ]
    for idx, (label, value, note) in enumerate(summary_items):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "EAF5EF" if idx < 2 else "FFF4D6")
        set_cell_border(cell, color="C9D3DF", size="6")
        set_cell_margins(cell, top=170, bottom=170, start=150, end=150)
        fill_cell(cell, label, size=8.5, bold=True, color=COLORS["muted"], align=WD_ALIGN_PARAGRAPH.CENTER)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(3)
        set_run(p2.add_run(value), size=14, bold=True, color=COLORS["navy"])
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        set_run(p3.add_run(note), size=8, color=COLORS["muted"])

    add_callout(
        doc,
        "CEO Takeaway",
        f"The company is currently committed to {money_php(current_committed_monthly)} per month. A realistic commercialization budget is {money_php(future_base_monthly)} per month. Activate the higher budget only when event volume and funding justify it.",
        fill="EAF5EF",
    )

    add_heading(doc, "What We Pay Every Month Today", 1)
    add_table(
        doc,
        ["Business Purpose", "What It Covers", "Monthly Cost"],
        [
            ("Website and application hosting", "Vercel Pro and Speed Insights", money_php(php(30))),
            ("Production database", "Supabase Pro for the live application", money_php(php(25))),
            ("UAT testing database", "Separate Supabase compute for safe pre-release testing", money_php(php(10))),
            ("Development assistants", "ChatGPT/Codex, Claude, and GitHub Copilot", money_php(current_ai)),
            ("TOTAL", "Current recurring subscriptions", money_php(current_committed_monthly)),
        ],
        [2600, 4400, 1960],
        font_size=9.0,
        right_cols=(2,),
    )

    add_heading(doc, "What Has Already Been Paid", 1)
    add_table(
        doc,
        ["Month", "What Was Purchased", "Amount"],
        [
            ("May 2026", "GitHub Copilot", money_php(php(39))),
            ("June 2026", "Hosting, Production/UAT databases, and development assistants", money_php(current_committed_monthly)),
            ("TOTAL", "Recorded software purchases since May", money_php(historical_total)),
        ],
        [1800, 5160, 2000],
        font_size=9.0,
        right_cols=(2,),
    )
    add_para(doc, f"Annualized current cost: approximately {money_php(current_committed_monthly * 12 + namecheap_recommended_first_year)}, including the current annual domain and DNS cost.", size=9.5, color=COLORS["muted"], italic=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Potential Monthly Cost", 1)
    add_para(doc, "Expected budget when Axon supports upcoming events with separate UAT and Production environments, stronger operating tools, and higher-capacity development support.", size=10.5, after=8)
    add_table(
        doc,
        ["Cost Group", "Why It Is Needed", "Monthly Budget"],
        [
            ("Core platform and two environments", "Hosting, Production and UAT databases, email, media, monitoring, security, and reservation capacity", money_php(runtime_base)),
            ("Commercial development capacity", "ChatGPT Pro and Claude Max for releases, event preparation, debugging, and urgent fixes", money_php(ai_future)),
            ("TOTAL", "Recommended funded commercialization budget", money_php(future_base_monthly)),
        ],
        [2500, 4460, 2000],
        font_size=9.0,
        right_cols=(2,),
    )
    add_callout(
        doc,
        "Why UAT and Production Are Separate",
        "UAT is a safe place to test changes before customers use them. Production remains stable for ticket buyers and event staff. This reduces the chance of registration, payment, email, QR ticket, or check-in changes affecting a live event.",
        fill="E8EEF5",
    )
    add_callout(
        doc,
        "Why Higher AI Plans Are Proposed",
        "During commercialization, UAT fixes and Production support happen in parallel. ChatGPT Pro and Claude Max provide more working capacity during releases and incidents. They support engineering work but do not replace human approval, QA, monitoring, or backups.",
        fill="E8EEF5",
    )
    add_heading(doc, "CEO Decisions and Cost Controls", 1)
    add_table(
        doc,
        ["Decision", "Recommended Treatment", "Financial Effect"],
        [
            ("GitHub Copilot", "Discontinue after transition to ChatGPT Pro and Claude Max", "Save PHP 2,262/month"),
            ("ChatGPT Pro + Claude Max", "Activate for commercialization and review after each event cycle", "PHP 12,290/month"),
            ("Google Maps paid capacity", "Keep as an optional reserve until usage requires it", "Up to PHP 15,950/month"),
            ("Legal, payment, security, emergency support", "Obtain quotations before budget approval", "Not yet included"),
        ],
        [2500, 4360, 2100],
        font_size=8.8,
        right_cols=(2,),
    )
    add_callout(
        doc,
        "Funding Presentation",
        f"Present {money_php(current_committed_monthly)}/month as today's recurring cost and {money_php(future_base_monthly)}/month as the funded commercialization target. Keep optional reserves and unquoted professional services separate.",
        fill="FFF4D6",
    )
    add_para(doc, f"Projected annual commercialization baseline: {money_php(future_base_monthly * 12 + hostinger_domain_first_year)}. Optional reserve scenario: {money_php(future_cons_monthly)}/month.", size=9.5, color=COLORS["muted"], italic=True)

    for current_section in doc.sections:
        footer = current_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run(footer.add_run("Axon Tickets CEO Cost Summary"), size=8.5, color=COLORS["muted"])
    return doc

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_metric_strip(doc)

    add_callout(
        doc,
        "CTO Recommendation",
        f"Report the current committed run-rate as USD {current_committed_usd} ({money_php(current_committed_monthly)}) per month. Separately disclose USD {historical_total_usd} ({money_php(historical_total)}) in recorded software expenses since May 2026.",
        fill="EAF5EF",
    )

    add_heading(doc, "Monthly and Annual Expense Summary", 1)
    add_kpi_table(doc)
    add_para(
        doc,
        "Monthly figures exclude annual domain charges. Annual totals add the selected domain scenario: current costs use Namecheap with PremiumDNS; projected scenarios use Hostinger first-year pricing.",
        size=9.5,
        color=COLORS["muted"],
        italic=True,
        after=6,
    )

    add_heading(doc, "Recorded Expenses Since May 2026", 1)
    add_historical_expense_table(doc)
    add_para(
        doc,
        f"May includes GitHub Copilot only (USD {historical_may_ai_usd}). June includes USD {historical_june_ai_usd} in AI subscriptions plus Vercel Pro, Supabase Pro, and the Axon UAT compute allocation totaling USD {historical_june_runtime_usd}. The cumulative recorded total is USD {historical_total_usd} ({money_php(historical_total)}).",
        size=9.5,
        color=COLORS["muted"],
        italic=True,
        after=6,
    )

    add_heading(doc, "Fixed Runtime Subscriptions", 1)
    runtime_rows = []
    for name, current, future, note in runtime:
        if isinstance(future, str):
            future_text = future
        elif future is None:
            future_text = "Pay-as-you-go"
        elif future == 0:
            future_text = "Free"
        else:
            future_text = money_php(future) + "/month"
        runtime_rows.append((name, current, future_text, note))
    add_table(doc, ["Tool", "Current", "Future", "CEO-Friendly Notes"], runtime_rows, [1550, 1050, 1650, 5110], font_size=8.2)

    add_heading(doc, "Developer and AI Productivity Tools", 1)
    add_table(doc, ["Tool", "Current", "Future", "CEO-Friendly Notes"], ai_tools, [1550, 1300, 1650, 4860], font_size=8.3)
    add_callout(
        doc,
        "AI Tooling Decision",
        f"Discontinue GitHub Copilot after transition. For commercialization, budget ChatGPT Pro at PHP 6,490/month and Claude Max at USD 100 ({money_php(php(100))})/month, totaling {money_php(ai_future)}/month. These are operational capacity upgrades, not hosting requirements.",
        fill="E8EEF5",
    )
    add_heading(doc, "Why Pro and Max Support UAT + Production", 2)
    add_bullets(
        doc,
        [
            "Two environments create parallel work: UAT validation and defect remediation must continue while Production remains stable and supportable.",
            "Higher usage limits reduce the risk of AI capacity being exhausted during release hardening, event-week changes, security reviews, and urgent incident response.",
            "ChatGPT Pro is the primary implementation and research workspace; Claude Max provides a second high-capacity review path for architecture, debugging, and independent verification.",
            "The upgrades should be activated for commercialization and reviewed after each event cycle. They improve engineering throughput but do not replace QA, approvals, monitoring, backups, or human accountability.",
        ],
    )

    # New page for assumptions and contingency.
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Domain and Planning Assumptions", 1)
    assumption_rows = [
        ("Current domain", "Namecheap: PHP 68/year (USD 1.18)", "Current registrar for axontickets.online."),
        ("Namecheap DNS add-on", "PremiumDNS: PHP 283/year (USD 4.88)", "Recommended if staying with Namecheap for stronger DNS uptime, DNSSEC, Anycast DNS, DNS-layer DDoS protection, and ALIAS record support."),
        ("Future domain option", f"Hostinger: {money_php(hostinger_domain_first_year)} first year", f"Succeeding-year renewal estimate: {money_php(hostinger_domain_renewal)}/year. Consider if it improves billing, support, DNS management, or ownership control."),
        ("Exchange rate", f"USD 1 = PHP {RATE}", f"Planning assumption as of {AS_OF_DATE}. USD vendor prices are authoritative; update only the PHP conversion when the rate changes."),
    ]
    add_table(doc, ["Item", "Amount", "Notes"], assumption_rows, [2200, 2200, 4960], font_size=8.8)

    add_heading(doc, "Commercialization Contingency Items", 1)
    add_para(
        doc,
        "These items are intentionally kept out of the fixed monthly and annual totals because pricing depends on supplier quotations, legal scope, payment gateway terms, actual usage, and final launch assumptions.",
        size=10.5,
        after=8,
    )
    contingency_rows = [
        ("Commercial Launch Contingency", "Legal documents, payment gateway onboarding, cybersecurity review, production readiness testing", "To be quoted"),
        ("Monthly Operations Contingency", "Monitoring, support tooling, backup retention, accounting, tax support", "To be quoted"),
        ("Emergency Application Contingency", "Urgent fixes, outage recovery, failed payment incidents, security response, emergency consultant support", "Reserved fund; amount to be approved"),
    ]
    add_table(doc, ["Group", "Includes", "Budget Status"], contingency_rows, [2300, 4760, 2300], header_fill="FFF4D6", font_size=8.6)

    add_heading(doc, "Contingency Scope Checklist", 2)
    add_bullets(
        doc,
        [
            "Payment and finance: BDO setup, monthly fee, transaction fee, refunds, chargebacks, accounting, tax, settlement, and PayMongo decision.",
            "Security: cybersecurity audit, penetration testing, vulnerability monitoring, secrets handling, admin access review, and incident response.",
            "Operations: backups, restore testing, disaster recovery, uptime monitoring, status page, alerts, runbook, admin manual, and support workflow.",
            "Legal readiness: terms of service, privacy policy, refund policy, organizer agreement, data retention, cookie/tracking disclosure, and payment agreement review.",
            "Product readiness: email deliverability, SPF/DKIM/DMARC, QA/UAT, load testing, mobile testing, admin permissions, Sentry ownership, and KPI dashboard definition.",
        ],
    )

    add_heading(doc, "Assumptions Still Needed to Finalize Budget", 2)
    add_bullets(
        doc,
        [
            "Expected events per month, attendees per event, ticket transactions, email volume, image/storage volume, Redis command volume, and support volume.",
            "Target commercial launch date, final BDO/payment gateway pricing, and final peso-to-USD exchange-rate assumption.",
            "Decision on whether hCaptcha Plus and Google Maps paid upgrade are needed immediately or should remain future reserves.",
        ],
    )

    add_callout(
        doc,
        "How to Use This Brief",
        f"Lead with the current committed cost: USD {current_committed_usd} ({money_php(current_committed_monthly)}) per month. Also disclose recorded software spending since May of USD {historical_total_usd} ({money_php(historical_total)}). Present projections and reserves separately.",
        fill="F2F4F7",
    )

    # Footer.
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_before = Pt(0)
        run = footer.add_run("Axon Tickets Commercialization Cost Brief")
        set_run(run, size=8.5, color=COLORS["muted"])

    return doc


if __name__ == "__main__":
    out = "reports/cto-cost-brief/Axon_Tickets_Commercialization_Cost_Brief.docx"
    doc = build_doc()
    doc.save(out)
    print(out)
